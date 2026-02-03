from typing import Optional, List, Set
from io import BytesIO

try:
    from docx import Document
except ImportError:
    Document = None


class TPDocumentConverter:
    @staticmethod
    def docx_to_markdown(file_path: Optional[str] = None, file_obj: Optional[BytesIO] = None) -> str:
        if Document is None:
            raise ValueError(
                "Модуль python-docx не установлен. "
                "Установите: poetry add python-docx"
            )
        try:
            if file_obj:
                file_obj.seek(0)
                doc = Document(file_obj)
            elif file_path:
                doc = Document(file_path)
            else:
                raise ValueError("Не указан ни путь к файлу, ни файловый объект")
            
            markdown_parts = []
            
            # Извлекаем header и footer из всех секций документа
            header_footer_parts = TPDocumentConverter._extract_header_footer(doc)
            if header_footer_parts:
                markdown_parts.extend(header_footer_parts)
            
            # Извлекаем основной контент документа
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    paragraph = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            paragraph = p
                            break
                    if paragraph and paragraph.text.strip():
                        markdown_text = TPDocumentConverter._paragraph_to_markdown(paragraph)
                        if markdown_text:
                            markdown_parts.append(markdown_text)
                elif element.tag.endswith('tbl'):
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    if table:
                        markdown_table = TPDocumentConverter._table_to_markdown(table)
                        if markdown_table:
                            markdown_parts.append(markdown_table)
            
            result = "\n\n".join(markdown_parts)
            if not result.strip():
                raise ValueError("Документ пуст или не удалось извлечь контент")
            return result.strip()
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Ошибка конвертации DOCX в Markdown: {str(e)}") from e
    
    @staticmethod
    def _extract_header_footer(doc) -> List[str]:
        """Извлекает текст из header и footer всех секций документа."""
        header_footer_parts = []
        seen_texts: Set[str] = set()
        
        for section in doc.sections:
            # Обрабатываем header
            if section.header:
                header_texts = []
                for paragraph in section.header.paragraphs:
                    text = paragraph.text.strip()
                    if text and text not in seen_texts:
                        header_texts.append(text)
                        seen_texts.add(text)
                if header_texts:
                    header_footer_parts.append("**Header:** " + " | ".join(header_texts))
                
                # Обрабатываем таблицы в header
                for table in section.header.tables:
                    table_text = TPDocumentConverter._table_to_markdown(table)
                    if table_text:
                        header_footer_parts.append("**Header Table:**\n" + table_text)
            
            # Обрабатываем footer
            if section.footer:
                footer_texts = []
                for paragraph in section.footer.paragraphs:
                    text = paragraph.text.strip()
                    if text and text not in seen_texts:
                        footer_texts.append(text)
                        seen_texts.add(text)
                if footer_texts:
                    header_footer_parts.append("**Footer:** " + " | ".join(footer_texts))
                
                # Обрабатываем таблицы в footer
                for table in section.footer.tables:
                    table_text = TPDocumentConverter._table_to_markdown(table)
                    if table_text:
                        header_footer_parts.append("**Footer Table:**\n" + table_text)
        
        return header_footer_parts

    @staticmethod
    def _paragraph_to_markdown(paragraph) -> str:
        text = paragraph.text.strip()
        if not text:
            return ""
        style = paragraph.style.name if paragraph.style else None
        if style and style.startswith('Heading'):
            level = 1
            if 'Heading 2' in style:
                level = 2
            elif 'Heading 3' in style:
                level = 3
            elif 'Heading 4' in style:
                level = 4
            elif 'Heading 5' in style:
                level = 5
            elif 'Heading 6' in style:
                level = 6
            return f"{'#' * level} {text}"
        return TPDocumentConverter._format_runs(paragraph.runs)

    @staticmethod
    def _format_runs(runs) -> str:
        result = []
        for run in runs:
            text = run.text
            if not text:
                continue
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"<u>{text}</u>"
            result.append(text)
        return "".join(result)

    @staticmethod
    def _table_to_markdown(table) -> str:
        if not table.rows:
            return ""
        markdown_rows = []
        for row_idx, row in enumerate(table.rows):
            cell_texts = []
            
            # Собираем текст из всех ячеек строки
            for cell in row.cells:
                cell_text = cell.text.strip()
                cell_text = " ".join(cell_text.split())
                cell_texts.append(cell_text)
            
            # Удаляем дубликаты: оставляем только первое вхождение, остальные очищаем
            seen_texts: Set[str] = set()
            cleaned_cells = []
            for cell_text in cell_texts:
                # Нормализуем текст для сравнения (приводим к нижнему регистру и удаляем лишние пробелы)
                normalized_text = " ".join(cell_text.lower().split()) if cell_text else ""
                
                if normalized_text and normalized_text in seen_texts:
                    # Дубликат - очищаем ячейку
                    cleaned_cells.append(" ")
                else:
                    # Уникальный текст - сохраняем
                    if normalized_text:
                        seen_texts.add(normalized_text)
                    # Экранируем специальные символы для markdown
                    escaped_text = cell_text.replace("|", "\\|").replace("\n", " ")
                    cleaned_cells.append(escaped_text or " ")
            
            markdown_row = "| " + " | ".join(cleaned_cells) + " |"
            markdown_rows.append(markdown_row)
            if row_idx == 0:
                separator = "| " + " | ".join(["---"] * len(cleaned_cells)) + " |"
                markdown_rows.append(separator)
        return "\n".join(markdown_rows)

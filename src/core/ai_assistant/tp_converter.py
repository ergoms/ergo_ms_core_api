"""
Сервис для конвертации документов техпроцессов (DOCX) в Markdown
"""
import logging
from pathlib import Path
from typing import Optional
from io import BytesIO

logger = logging.getLogger(__name__)


class TPDocumentConverter:
    """
    Конвертер документов техпроцессов из DOCX в Markdown
    
    Выполняет кастомную обработку таблиц с сохранением структуры,
    форматирования и метаданных.
    """
    
    @staticmethod
    def docx_to_markdown(file_path: Optional[str] = None, file_obj: Optional[BytesIO] = None) -> str:
        """
        Конвертирует DOCX файл в Markdown
        
        Args:
            file_path: Путь к файлу на диске
            file_obj: Файловый объект (BytesIO)
            
        Returns:
            Markdown контент документа
            
        Raises:
            ValueError: При ошибке конвертации
        """
        try:
            from docx import Document
        except ImportError:
            raise ValueError(
                "Модуль python-docx не установлен. "
                "Установите: poetry add python-docx"
            )
        
        try:
            if file_obj:
                file_obj.seek(0)
                doc = Document(file_obj)
            elif file_path:
                doc = Document(file_path)
            else:
                raise ValueError("Не указан ни путь к файлу, ни файловый объект")
            
            markdown_parts = []
            
            # Обрабатываем все элементы документа по порядку
            for element in doc.element.body:
                # Параграфы
                if element.tag.endswith('p'):
                    paragraph = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            paragraph = p
                            break
                    
                    if paragraph and paragraph.text.strip():
                        markdown_text = TPDocumentConverter._paragraph_to_markdown(paragraph)
                        if markdown_text:
                            markdown_parts.append(markdown_text)
                
                # Таблицы
                elif element.tag.endswith('tbl'):
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    
                    if table:
                        markdown_table = TPDocumentConverter._table_to_markdown(table)
                        if markdown_table:
                            markdown_parts.append(markdown_table)
            
            result = "\n\n".join(markdown_parts)
            
            if not result.strip():
                raise ValueError("Документ пуст или не удалось извлечь контент")
            
            return result.strip()
            
        except Exception as e:
            if isinstance(e, ValueError):
                raise
            raise ValueError(f"Ошибка конвертации DOCX в Markdown: {str(e)}") from e
    
    @staticmethod
    def _paragraph_to_markdown(paragraph) -> str:
        """
        Конвертирует параграф в Markdown с сохранением форматирования
        """
        text = paragraph.text.strip()
        if not text:
            return ""
        
        # Проверяем стиль параграфа
        style = paragraph.style.name if paragraph.style else None
        
        # Заголовки
        if style and style.startswith('Heading'):
            level = 1
            if 'Heading 1' in style:
                level = 1
            elif 'Heading 2' in style:
                level = 2
            elif 'Heading 3' in style:
                level = 3
            elif 'Heading 4' in style:
                level = 4
            elif 'Heading 5' in style:
                level = 5
            elif 'Heading 6' in style:
                level = 6
            
            return f"{'#' * level} {text}"
        
        # Обычный текст с форматированием
        formatted_text = TPDocumentConverter._format_runs(paragraph.runs)
        
        return formatted_text
    
    @staticmethod
    def _format_runs(runs) -> str:
        """
        Форматирует runs (части текста) с сохранением жирного, курсива и т.д.
        """
        result = []
        for run in runs:
            text = run.text
            if not text:
                continue
            
            # Жирный
            if run.bold:
                text = f"**{text}**"
            
            # Курсив
            if run.italic:
                text = f"*{text}*"
            
            # Подчеркнутый (в Markdown нет нативного подчеркивания, используем <u>)
            if run.underline:
                text = f"<u>{text}</u>"
            
            result.append(text)
        
        return "".join(result)
    
    @staticmethod
    def _table_to_markdown(table) -> str:
        """
        Конвертирует таблицу в Markdown таблицу с сохранением структуры
        
        Обрабатывает таблицы как в примере скрипта - извлекает все ячейки
        и формирует Markdown таблицу.
        """
        if not table.rows:
            return ""
        
        markdown_rows = []
        
        # Обрабатываем каждую строку
        for row_idx, row in enumerate(table.rows):
            cells = []
            
            for cell in row.cells:
                # Извлекаем текст из ячейки (может быть многострочным)
                cell_text = cell.text.strip()
                # Заменяем переносы строк на пробелы для Markdown таблицы
                cell_text = " ".join(cell_text.split())
                # Экранируем символы, которые могут сломать таблицу
                cell_text = cell_text.replace("|", "\\|").replace("\n", " ")
                cells.append(cell_text or " ")
            
            # Формируем строку таблицы
            markdown_row = "| " + " | ".join(cells) + " |"
            markdown_rows.append(markdown_row)
            
            # После первой строки добавляем разделитель
            if row_idx == 0:
                separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                markdown_rows.append(separator)
        
        return "\n".join(markdown_rows)

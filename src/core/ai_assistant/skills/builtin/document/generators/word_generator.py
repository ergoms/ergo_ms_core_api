"""
Генератор Word документов.
Минимальная версия без сложных элементов.
"""
import logging
from pathlib import Path
from typing import List, Dict, Any

from docx import Document

logger = logging.getLogger(__name__)


class WordGenerator:
    """Генератор Word документов."""
    
    def generate_bi_report(
        self,
        output_path: Path,
        title: str,
        date: str,
        file_name: str,
        question: str,
        commentary: str,
        data: List[Dict[str, Any]],
        columns: List[str],
        sql: str
    ) -> Path:
        """Генерирует Word отчёт по BI анализу."""
        
        logger.info(f"Генерация BI отчёта: {title}")
        
        # Создаем документ
        doc = Document()
        
        # Заголовок - простой текст
        doc.add_paragraph(self._safe(title))
        doc.add_paragraph("")
        
        # Метаинформация
        doc.add_paragraph(f"Дата: {self._safe(date)}")
        doc.add_paragraph(f"Файл: {self._safe(file_name)}")
        doc.add_paragraph(f"Запрос: {self._safe(question)}")
        doc.add_paragraph("")
        
        # Анализ
        doc.add_paragraph("АНАЛИЗ ДАННЫХ")
        if commentary:
            doc.add_paragraph(self._safe(commentary))
        doc.add_paragraph("")
        
        # Данные - просто текстом
        doc.add_paragraph("РЕЗУЛЬТАТЫ")
        if data and columns:
            # Заголовки колонок
            doc.add_paragraph("Колонки: " + ", ".join(self._safe(str(c)) for c in columns))
            
            # Данные (максимум 10 строк)
            for i, row in enumerate(data[:10]):
                if isinstance(row, dict):
                    row_text = " | ".join(f"{c}: {self._safe(str(row.get(c, '')))}" for c in columns)
                else:
                    row_text = str(row)
                doc.add_paragraph(f"  {i+1}. {row_text}")
            
            if len(data) > 10:
                doc.add_paragraph(f"  ... и ещё {len(data) - 10} строк")
        else:
            doc.add_paragraph("Нет данных")
        
        doc.add_paragraph("")
        
        # SQL
        doc.add_paragraph("SQL ЗАПРОС")
        if sql:
            doc.add_paragraph(self._safe(sql))
        
        doc.add_paragraph("")
        doc.add_paragraph("Документ создан системой ERGO MS")
        
        # Сохраняем
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"Сохраняем документ: {output_path}")
        doc.save(str(output_path))
        
        if output_path.exists():
            size = output_path.stat().st_size
            logger.info(f"Документ создан, размер: {size} байт")
        
        return output_path
    
    def generate_simple(
        self,
        output_path: Path,
        title: str,
        content: str,
        author: str = "AI",
        date: str = ""
    ) -> Path:
        """Генерирует простой документ."""
        
        doc = Document()
        
        doc.add_paragraph(self._safe(title))
        doc.add_paragraph("")
        
        if date:
            doc.add_paragraph(f"Дата: {date}")
        doc.add_paragraph(f"Автор: {author}")
        doc.add_paragraph("")
        
        if content:
            doc.add_paragraph(self._safe(content))
        
        doc.add_paragraph("")
        doc.add_paragraph("ERGO MS")
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return output_path
    
    def _safe(self, text) -> str:
        """Очищает текст."""
        if text is None:
            return ""
        
        text = str(text)
        
        # Удаляем управляющие символы
        result = []
        for char in text:
            code = ord(char)
            if code >= 32 or char in '\n\r\t':
                if not (0xD800 <= code <= 0xDFFF):
                    if code != 0xFFFE and code != 0xFFFF:
                        result.append(char)
        
        return ''.join(result)
